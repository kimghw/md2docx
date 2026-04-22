#!/usr/bin/env python3
"""C9 — inject extracted table styles into a blank docx and render both.

Inputs:
  - extract_out_dir  : dir with table_style_bundle/styles_excerpt.xml +
                       sample_table.xml (from extract.py)
  - out_dir          : where injected.docx + its PDF/PNG go

What it does:
  1) Build a blank docx with a single 3x3 plain table (python-docx).
  2) Inject all <w:style> blocks from styles_excerpt.xml into its
     word/styles.xml (replace if the styleId already exists, else append
     before </w:styles>).
  3) Set the table's <w:tblPr> to reference the anchor styleId and copy
     the original <w:tblLook> so conditional formatting (firstRow, firstCol)
     actually fires.
  4) Render the injected docx to PDF + PNG via docx2pdf + PyMuPDF.
  5) Also render the source fixture to PDF + PNG for side-by-side baseline.

The baseline and injected PNGs can then be compared (visually, by an agent
or a human). If the two tables look the same, the extracted bundle is
functionally reusable.
"""
from __future__ import annotations

import os
import re
import shutil
import sys
import zipfile
from typing import Optional

try:
    from docx import Document
    from docx2pdf import convert as docx2pdf_convert
    import fitz  # PyMuPDF
except ModuleNotFoundError as e:
    sys.exit(f"missing module: {e}")

HERE = os.path.dirname(os.path.abspath(__file__))


# ---- XML helpers ----
_STYLE_BLOCK = re.compile(r'<w:style\b[^>]*>.*?</w:style>', re.S)
_STYLEID_ATTR = re.compile(r'\bw:styleId="([^"]*)"')


def style_blocks_with_ids(xml_text: str) -> list[tuple[str, str]]:
    """Return [(styleId, full <w:style>…</w:style> block), …]."""
    out = []
    for m in _STYLE_BLOCK.finditer(xml_text):
        block = m.group(0)
        sid_m = _STYLEID_ATTR.search(block)
        out.append((sid_m.group(1) if sid_m else "", block))
    return out


def read_zip(zp: str) -> dict[str, bytes]:
    with zipfile.ZipFile(zp) as z:
        return {n: z.read(n) for n in z.namelist()}


def write_zip(zp: str, members: dict[str, bytes]) -> None:
    with zipfile.ZipFile(zp, "w", zipfile.ZIP_DEFLATED) as z:
        for n, data in members.items():
            z.writestr(n, data)


# ---- build blank docx ----
def build_blank_docx(path: str) -> None:
    d = Document()
    d.add_heading("injected-style test", level=1)
    d.add_paragraph("Below: a 3x3 plain table (style is injected post-hoc).")
    t = d.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t.cell(r, c).text = f"R{r}C{c}"
    d.save(path)


# ---- core: inject + restyle ----
def read_anchor_and_tbllook(sample_table_xml: str) -> tuple[Optional[str], Optional[str]]:
    """Extract anchor styleId and the <w:tblLook .../> element."""
    anchor = None
    m = re.search(r'<w:tblStyle\s+w:val="([^"]*)"', sample_table_xml)
    if m:
        anchor = m.group(1)
    tbllook = None
    m = re.search(r'<w:tblLook\b[^/]*/>', sample_table_xml)
    if m:
        tbllook = m.group(0)
    return anchor, tbllook


def inject_styles(target_styles_xml: str, donor_blocks: list[tuple[str, str]]) -> str:
    """Insert donor <w:style> blocks into target_styles_xml.

    If a styleId already exists in target, replace its block; otherwise
    insert before </w:styles>.
    """
    out = target_styles_xml
    for sid, block in donor_blocks:
        if not sid:
            continue
        existing = re.search(
            r'<w:style\b[^>]*w:styleId="' + re.escape(sid) + r'"[^>]*>.*?</w:style>',
            out, re.S)
        if existing:
            out = out[:existing.start()] + block + out[existing.end():]
        else:
            out = out.replace("</w:styles>", block + "</w:styles>", 1)
    return out


def restyle_first_table(doc_xml: str, anchor: str, tbllook: Optional[str]) -> str:
    """Rewrite the first <w:tbl>'s <w:tblPr> to reference anchor+tblLook."""
    m = re.search(r'(<w:tbl\b[^>]*>)(.*?)(</w:tbl>)', doc_xml, re.S)
    if not m:
        raise RuntimeError("no <w:tbl> in target document.xml")
    body = m.group(2)
    # pull existing tblPr (if any), or insert a fresh one
    tblpr_m = re.search(r'<w:tblPr\b[^>]*>.*?</w:tblPr>|<w:tblPr\b[^/]*/>',
                        body, re.S)
    new_tblpr_parts = [f'<w:tblStyle w:val="{anchor}"/>',
                       '<w:tblW w:w="0" w:type="auto"/>']
    if tbllook:
        new_tblpr_parts.append(tbllook)
    new_tblpr = "<w:tblPr>" + "".join(new_tblpr_parts) + "</w:tblPr>"
    if tblpr_m:
        body_new = body[:tblpr_m.start()] + new_tblpr + body[tblpr_m.end():]
    else:
        body_new = new_tblpr + body
    return doc_xml[:m.start()] + m.group(1) + body_new + m.group(3) + doc_xml[m.end():]


def inject_bundle_into_docx(blank_docx: str, out_docx: str,
                            styles_excerpt_path: str,
                            sample_table_path: str) -> dict:
    """Produce out_docx = blank_docx + injected table style + retyped table."""
    excerpt = open(styles_excerpt_path, encoding="utf-8").read()
    sample = open(sample_table_path, encoding="utf-8").read()
    anchor, tbllook = read_anchor_and_tbllook(sample)
    if not anchor:
        raise RuntimeError("no anchor styleId in sample_table.xml")

    donor_blocks = style_blocks_with_ids(excerpt)

    members = read_zip(blank_docx)
    styles_xml = members["word/styles.xml"].decode("utf-8")
    doc_xml = members["word/document.xml"].decode("utf-8")

    styles_xml = inject_styles(styles_xml, donor_blocks)
    doc_xml = restyle_first_table(doc_xml, anchor, tbllook)

    members["word/styles.xml"] = styles_xml.encode("utf-8")
    members["word/document.xml"] = doc_xml.encode("utf-8")
    write_zip(out_docx, members)

    return {
        "anchor": anchor,
        "injected_style_ids": [sid for sid, _ in donor_blocks if sid],
        "tbllook_copied": tbllook is not None,
    }


# ---- render ----
def render_docx(src_docx: str, out_pdf: str, out_dir: str, prefix: str) -> list[str]:
    docx2pdf_convert(os.path.abspath(src_docx), os.path.abspath(out_pdf))
    doc = fitz.open(out_pdf)
    pages = doc.page_count
    pngs = []
    try:
        zoom = 200 / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        for i in range(pages):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            p = os.path.join(out_dir, f"{prefix}.page-{i + 1:02d}.png")
            pix.save(p)
            pngs.append(p)
    finally:
        doc.close()
    return pngs


# ---- verifier: structural sanity via python-docx reopen ----
def verify_structural(injected_docx: str) -> tuple[bool, str]:
    try:
        Document(injected_docx)
    except Exception as e:
        return False, f"python-docx failed to open: {e}"
    return True, "opens cleanly"


# ---- one-shot driver ----
def run_c9(extract_out: str, source_docx: str, out_dir: str) -> dict:
    os.makedirs(out_dir, exist_ok=True)
    blank = os.path.join(out_dir, "blank.docx")
    injected = os.path.join(out_dir, "injected.docx")

    build_blank_docx(blank)

    excerpt_path = os.path.join(extract_out, "table_style_bundle",
                                "styles_excerpt.xml")
    sample_path = os.path.join(extract_out, "sample_table.xml")
    meta = inject_bundle_into_docx(blank, injected, excerpt_path, sample_path)

    ok, note = verify_structural(injected)
    meta["structural_ok"] = ok
    meta["structural_note"] = note

    try:
        inj_pngs = render_docx(
            injected,
            os.path.join(out_dir, "injected.pdf"),
            out_dir, "injected")
        meta["injected_pngs"] = inj_pngs
    except Exception as e:
        meta["injected_pngs"] = []
        meta["render_error_injected"] = str(e)

    try:
        base_pngs = render_docx(
            source_docx,
            os.path.join(out_dir, "baseline.pdf"),
            out_dir, "baseline")
        meta["baseline_pngs"] = base_pngs
    except Exception as e:
        meta["baseline_pngs"] = []
        meta["render_error_baseline"] = str(e)

    return meta


def main():
    import argparse
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--extract-out", required=True,
                    help="Directory produced by extract.py")
    ap.add_argument("--source", required=True,
                    help="Original fixture docx (baseline render)")
    ap.add_argument("--out-dir", required=True,
                    help="Where injected.docx + PNGs go")
    args = ap.parse_args()

    result = run_c9(args.extract_out, args.source, args.out_dir)
    print("anchor             :", result["anchor"])
    print("injected styleIds  :", ", ".join(result["injected_style_ids"]))
    print("tblLook copied     :", result["tbllook_copied"])
    print("structural ok      :", result["structural_ok"], "-",
          result["structural_note"])
    print("baseline PNGs      :", len(result.get("baseline_pngs", [])))
    print("injected PNGs      :", len(result.get("injected_pngs", [])))
    for key in ("render_error_injected", "render_error_baseline"):
        if key in result:
            print(f"  {key}: {result[key]}")


if __name__ == "__main__":
    main()
