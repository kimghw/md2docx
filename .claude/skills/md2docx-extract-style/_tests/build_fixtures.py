#!/usr/bin/env python3
"""Build known-input fixtures for validating extract-docx-styles.

Each fixture targets a specific contract clause of the skill:
  fx_empty.docx        — no tables (② must be skipped)
  fx_tablegrid.docx    — single Table Grid (has_style_Table=false, chain resolves)
  fx_styled_Table.docx — same but styleId renamed to "Table" (has_style_Table=true)
  fx_multi.docx        — 3 tables (sample-index selection for table_sources.xml)
"""
from __future__ import annotations

import os
import re
import sys
import zipfile

try:
    from docx import Document
except ModuleNotFoundError:
    sys.exit("python-docx not installed: pip install python-docx")


HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "fixtures")


def _add_simple_table(doc, rows=3, cols=3, tag="T"):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    for r in range(rows):
        for c in range(cols):
            t.cell(r, c).text = f"{tag}{r}{c}"
    return t


def build_empty(path):
    d = Document()
    d.add_heading("제목 1", level=1)
    d.add_paragraph("일반 문단입니다.")
    d.save(path)


def build_tablegrid(path):
    d = Document()
    d.add_heading("제목 1", level=1)
    d.add_paragraph("아래는 단일 표.")
    _add_simple_table(d, tag="G")
    d.save(path)


def _rename_table_styleid(docx_path, new_id="Table"):
    """Rename whatever styleId the first <w:tbl> uses to `new_id`.

    python-docx normally gives "Table Grid" a styleId like "TableGrid".
    We overwrite both document.xml (the reference) and styles.xml (the definition)
    so that extract.py reports has_style_Table=true.
    """
    with zipfile.ZipFile(docx_path) as z:
        members = {n: z.read(n) for n in z.namelist()}
    doc_xml = members["word/document.xml"].decode("utf-8")
    styles_xml = members["word/styles.xml"].decode("utf-8")

    m = re.search(r'<w:tblStyle\s+w:val="([^"]*)"', doc_xml)
    if not m:
        raise RuntimeError("no <w:tblStyle> in document.xml — nothing to rename")
    old_id = m.group(1)

    doc_xml = re.sub(
        r'<w:tblStyle\s+w:val="' + re.escape(old_id) + r'"',
        f'<w:tblStyle w:val="{new_id}"',
        doc_xml,
    )
    styles_xml = re.sub(
        r'(<w:style\b[^>]*w:styleId=")' + re.escape(old_id) + r'"',
        r'\1' + new_id + '"',
        styles_xml,
    )
    members["word/document.xml"] = doc_xml.encode("utf-8")
    members["word/styles.xml"] = styles_xml.encode("utf-8")

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in members.items():
            z.writestr(name, data)


def build_styled_Table(path):
    build_tablegrid(path)
    _rename_table_styleid(path, new_id="Table")


def build_multi(path):
    d = Document()
    d.add_heading("Multi Table", level=1)
    for i in range(3):
        d.add_paragraph(f"Table {i} (index {i}):")
        _add_simple_table(d, rows=2, cols=2, tag=f"M{i}-")
    d.save(path)


def main():
    os.makedirs(OUT, exist_ok=True)
    builders = [
        ("fx_empty.docx", build_empty),
        ("fx_tablegrid.docx", build_tablegrid),
        ("fx_styled_Table.docx", build_styled_Table),
        ("fx_multi.docx", build_multi),
    ]
    for name, fn in builders:
        p = os.path.join(OUT, name)
        fn(p)
        print(f"  built  {os.path.relpath(p, HERE)}  ({os.path.getsize(p)} bytes)")


if __name__ == "__main__":
    main()
